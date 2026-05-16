from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
SCRIPT_PATH = ROOT / "presentation" / "pcos_pathfinder_script.md"
OUT_PATH = ROOT / "presentation" / "pcos_pathfinder_slide_script.docx"

ACCENT = RGBColor(46, 116, 181)
DARK_ACCENT = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)


def clean_text(text: str) -> str:
    replacements = {
        "\ufffd": "-",
        "\u2013": "-",
        "\u2014": "-",
        "\u2192": "->",
        "\u0394": "Delta",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    return text.strip()


def parse_slides(markdown: str) -> list[dict[str, str]]:
    lines = markdown.splitlines()
    slides: list[dict[str, str]] = []
    current: dict[str, str] | None = None
    body: list[str] = []

    heading_re = re.compile(r"^## Slide\s+(\d+)\s+.\s+(.+?)\s+\(([^)]*)\)(?:\s+.\s+.*)?$")

    for line in lines:
        match = heading_re.match(line)
        if match:
            if current is not None:
                current["script"] = clean_text("\n".join(body))
                slides.append(current)
            number, title, time = match.groups()
            current = {
                "number": number,
                "title": clean_text(title),
                "time": clean_text(time),
                "script": "",
            }
            body = []
            continue

        if current is None:
            continue
        if line.startswith("---") or line.startswith("## Speaker tips"):
            break
        if line.strip():
            body.append(line)

    if current is not None:
        current["script"] = clean_text("\n".join(body))
        slides.append(current)

    return slides


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.space_after = Pt(10)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, DARK_ACCENT, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_slide_entry(doc: Document, slide: dict[str, str]) -> None:
    heading = doc.add_heading(
        f"Slide {slide['number']}: {slide['title']} ({slide['time']})",
        level=1,
    )
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_after = Pt(10)
    for idx, part in enumerate(slide["script"].split("\n")):
        if idx:
            p.add_run().add_break(WD_BREAK.LINE)
        p.add_run(part)


def build() -> None:
    markdown = SCRIPT_PATH.read_text(encoding="utf-8")
    slides = parse_slides(markdown)
    if not slides:
        raise RuntimeError(f"No slide scripts parsed from {SCRIPT_PATH}")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "PCOS Pathfinder - Slide Script"
    header.style = doc.styles["Normal"]
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Page ").font.size = Pt(9)
    add_page_number(footer)

    doc.add_paragraph("PCOS Pathfinder Slide Script", style="Title")
    subtitle = doc.add_paragraph(
        "Slide number : script reference, updated with the blood-group evidence note from Yang et al. 2025.",
        style="Normal",
    )
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.color.rgb = MUTED

    callout = doc.add_table(rows=1, cols=1)
    callout.autofit = False
    callout.columns[0].width = Inches(6.5)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, "E8EEF5")
    cell.paragraphs[0].text = (
        "Blood-group script justification: Yang et al. (BMC Endocrine Disorders, 2025) found no significant "
        "ABO distribution difference between PCOS and controls, but did find within-PCOS associations with "
        "menstrual bleeding level and BMI, FSH, LH, and estradiol. Therefore, ABO/Rh should be treated as "
        "categorical symptom-severity metadata, not as the numeric 11-18 diagnostic input."
    )
    cell.paragraphs[0].style = doc.styles["Normal"]
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    for slide in slides:
        add_slide_entry(doc, slide)

    doc.add_heading("Source", level=1)
    doc.add_paragraph(
        "Yang S, Zhang H, Shi L, et al. Menstrual disorder is associated with blood type in PCOS patients: "
        "evidence from a cross-sectional survey. BMC Endocrine Disorders. 2025;25:77. "
        "https://doi.org/10.1186/s12902-025-01898-0",
        style="Normal",
    )

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build()
